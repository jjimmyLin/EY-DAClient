from pathlib import Path
import sys

from docx import Document

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

paths = list((Path.home() / "Desktop").glob("*dify*docx"))
print([str(path) for path in paths])
path = paths[0]
document = Document(str(path))
print("PARAGRAPHS", len(document.paragraphs))
print("TABLES", len(document.tables))

for index, paragraph in enumerate(document.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        print(f"P{index}: {text}")

for table_index, table in enumerate(document.tables, start=1):
    print(f"TABLE {table_index}")
    for row_index, row in enumerate(table.rows, start=1):
        values = [
            cell.text.strip().replace("\n", " / ")
            for cell in row.cells
        ]
        print(f"R{row_index}: " + " || ".join(values))
