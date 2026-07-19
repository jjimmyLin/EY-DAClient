from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "EY_DataPilot_User_Guide.docx"


ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(32, 33, 36)
MUTED = RGBColor(95, 99, 104)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F7F9FC"
BORDER = "DADCE0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, value in values.items():
        tag = f"w:{side}"
        node = margins.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in_inches: list[float]) -> None:
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    widths_dxa = [int(width * 1440) for width in widths_in_inches]
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)


def set_paragraph_font(paragraph, font_name: str = "Microsoft YaHei", size: int | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if size:
            run.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("EY DataPilot 用户使用指南")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("面向业务用户的数据导入、清洗、智能分析与结果导出操作手册")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_cell_margins(meta)
    set_table_width(meta, [1.6, 4.9])
    rows = [
        ("适用版本", "当前桌面端版本"),
        ("适用对象", "审计、咨询、财务、运营及其他需要处理 Excel 数据的业务用户"),
        ("核心流程", "导入数据集 → 选择功能 → 数据清洗 / 数据分析 → 查看结果 → 导出"),
        ("当前限制", "支持 .xlsx / .xls / .xlsm；单文件最大 1GB；分析最多选择 3 个数据集；清洗一次处理 1 个数据集"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], BLUE_FILL)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
        for cell in row.cells:
            set_paragraph_font(cell.paragraphs[0])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(6)
    r = note.add_run("建议阅读方式：")
    r.bold = True
    r.font.color.rgb = DARK
    note.add_run(" 第一次使用请按第 3 章快速开始操作；日常使用时可直接查看第 5 至第 8 章。")
    set_paragraph_font(note)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table, top=110, bottom=110, start=160, end=160)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "：")
    r.bold = True
    r.font.color.rgb = DARK
    p.add_run(body)
    set_paragraph_font(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        set_paragraph_font(p)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        set_paragraph_font(p)


def add_feature_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, [1.35, 1.45, 2.55, 1.15])
    headers = ["模块", "用户动作", "系统输出", "适用场景"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE_FILL)
        p = cell.paragraphs[0]
        p.add_run(header).bold = True
        set_paragraph_font(p, size=9)
    rows = [
        ("数据导入", "拖拽或选择 Excel 文件", "缓存、识别工作表、生成基础数据概览", "开始任何任务前"),
        ("数据资料库", "打开 Datasets 面板并高亮数据集", "确定当前分析或清洗范围", "多文件管理"),
        ("数据清洗", "扫描后选择固定清洗规则", "生成新的清洗后 Excel 文件", "脏数据处理"),
        ("数据分析", "输入自然语言问题并应用代码", "答案、指标、表格、图表、洞察", "业务问答与探索"),
        ("结果导出", "选择全部或单个结果导出", "Excel 格式分析结果文件", "汇报、复核、留档"),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.add_run(value)
            set_paragraph_font(p, size=9)


def build_doc() -> Document:
    doc = Document()
    apply_styles(doc)
    add_title_block(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)
    set_paragraph_font(footer, size=9)

    doc.add_heading("1. 产品定位", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "EY DataPilot 是一款面向业务人员的桌面端数据处理与分析工具。"
        "它将数据导入、数据清洗、AI 辅助分析、本地 Python 执行、结构化结果展示和结果导出整合到一个工作流中。"
    )
    set_paragraph_font(p)
    add_callout(
        doc,
        "核心原则",
        "用户负责提出业务问题和选择处理方式；软件负责稳定执行、结构化展示和尽量保留可复核过程。",
    )

    doc.add_heading("2. 当前支持范围", level=1)
    add_feature_table(doc)
    doc.add_heading("2.1 文件与处理限制", level=2)
    add_bullets(
        doc,
        [
            "支持导入 Excel 工作簿：.xlsx、.xls、.xlsm。",
            "当前版本不支持直接导入 CSV；如需使用 CSV，请先转换为 Excel 工作簿。",
            "单个文件大小上限为 1GB。",
            "数据分析模式最多同时选择 3 个数据集。",
            "数据清洗模式一次只处理 1 个目标数据集。",
            "分析结果导出为 .xlsx 文件；清洗结果也会另存为新的 .xlsx 文件。",
        ],
    )

    doc.add_heading("3. 快速开始", level=1)
    add_numbers(
        doc,
        [
            "启动应用，停留在初始数据导入页面。",
            "点击导入框或将 Excel 文件拖入导入区域。",
            "等待数据集进入 Ready 状态。",
            "选择 Data Analyze 或 Data Clean。",
            "在右侧或顶部的 Datasets 数据资料库中高亮当前要处理的数据集。",
            "进入数据清洗时先 Scan，再勾选适用规则并执行；进入数据分析时输入问题，等待代码生成后点击 Apply。",
            "查看结构化结果，并按需导出。",
        ],
    )
    add_callout(
        doc,
        "推荐习惯",
        "先导入所有需要用到的数据，再选择功能。数据资料库是全局的，清洗和分析共享同一批已导入数据。",
    )

    doc.add_heading("4. 导入数据集", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "数据导入是使用软件的第一步。导入后，文件会被登记为应用内的数据集，并进入统一的数据资料库。"
        "后续无论使用数据清洗还是数据分析，都从这个资料库选择目标数据。"
    )
    set_paragraph_font(p)
    doc.add_heading("4.1 导入方式", level=2)
    add_bullets(
        doc,
        [
            "点击首页导入框，选择一个或多个 Excel 文件。",
            "将 Excel 文件直接拖拽到导入区域。",
            "如文件类型或大小不符合要求，系统会提示并停止导入。",
        ],
    )
    doc.add_heading("4.2 导入后的状态", level=2)
    add_bullets(
        doc,
        [
            "Ready：数据集已准备好，可以用于清洗或分析。",
            "Importing / Processing：系统正在读取和缓存数据，请等待。",
            "Failed：导入失败，可根据错误提示检查文件类型、大小或文件内容。",
        ],
    )

    doc.add_heading("5. 使用 Data Library", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Data Library 是全局数据资料库。进入功能页后，可以通过顶部的 Datasets 按钮打开。"
        "点击界面其他区域时，资料库会自动关闭。"
    )
    set_paragraph_font(p)
    add_bullets(
        doc,
        [
            "在数据分析模式中，高亮 1 到 3 个数据集作为分析范围。",
            "在数据清洗模式中，只能高亮 1 个数据集作为清洗目标。",
            "可以查看数据集 overview，帮助快速理解数据内容。",
            "可以删除不再需要的数据集；删除只会移除应用中的数据记录，不代表覆盖或修改原始文件。",
        ],
    )

    doc.add_heading("6. 数据清洗操作", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "数据清洗功能采用固定、可控、可配置的处理逻辑。软件会先扫描数据，识别当前数据中实际存在的问题，"
        "再启用适用的清洗规则；不适用的规则会保持禁用，避免用户误操作。"
    )
    set_paragraph_font(p)
    doc.add_heading("6.1 基本流程", level=2)
    add_numbers(
        doc,
        [
            "进入 Data Clean。",
            "在 Data Library 中高亮一个目标数据集。",
            "点击 Scan，等待系统扫描脏数据问题。",
            "查看每项规则的状态、受影响列和预计修改数量。",
            "勾选需要执行的清洗规则；对于缺失值和混合数值等规则，可按列配置处理方式。",
            "点击 Execute，选择清洗后文件的保存路径。",
            "等待后台执行完成，打开新生成的清洗后文件进行复核。",
        ],
    )
    doc.add_heading("6.2 常见清洗能力", level=2)
    add_bullets(
        doc,
        [
            "缺失值处理：按列选择填充、删除或保持不变等方式。",
            "重复行处理：识别完全重复或关键字段重复记录。",
            "混合数值处理：区分真实 NULL 与混入数字列的文本异常值。",
            "非法 Excel 字符处理：修复无法安全写入 Excel 的字符。",
            "公式注入防护：防止以 =、+、-、@ 等开头的文本被 Excel 解释为公式。",
            "原始行号保留：便于追踪清洗结果对应的原始数据位置。",
        ],
    )
    add_callout(
        doc,
        "重要安全设计",
        "数据清洗不会覆盖原始文件。执行清洗时必须另存为新的 Excel 文件，以便保留原始数据用于复核。",
    )

    doc.add_heading("7. 数据分析操作", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "数据分析功能允许用户用自然语言提出问题。系统会将问题、数据结构和样本信息发送到 Dify 工作流，"
        "由 AI 生成 Python 代码；代码在本地软件中预检、展示、应用并执行，最终以结构化结果呈现。"
    )
    set_paragraph_font(p)
    doc.add_heading("7.1 基本流程", level=2)
    add_numbers(
        doc,
        [
            "进入 Data Analysis。",
            "在 Data Library 中高亮 1 到 3 个数据集。",
            "在输入框中描述分析需求，可以一次输入多个问题。",
            "等待 Dify 返回分析代码和分析计划。",
            "如代码通过代表性样本预检，系统默认展示 Python 页面。",
            "检查代码，如需调整可手动修改；点击 Reset 可恢复生成版本。",
            "点击 Python 页面右上角 Apply，运行完整数据分析。",
            "在 Result 页面查看结果。",
        ],
    )
    doc.add_heading("7.2 多问题结果", level=2)
    add_bullets(
        doc,
        [
            "如果一次输入多个需求，结果会拆成多个 result panel。",
            "每个 result panel 对应一个用户问题。",
            "每个问题可包含答案、支持指标、支持表格、支持图表和洞察。",
            "Python 代码仍然是一整段，便于统一审查和执行。",
        ],
    )
    doc.add_heading("7.3 如何写出更好的问题", level=2)
    add_bullets(
        doc,
        [
            "明确数据集：例如“使用销售数据和售后数据”。",
            "明确字段：例如“按车型”“按月份”“以 customer_id 关联”。",
            "明确指标：例如“总销售额”“平均维修成本”“重复到店次数”。",
            "明确输出：例如“给出排名表”“画趋势图”“只返回异常记录”。",
            "多文件分析时说明关联逻辑，避免 AI 猜测错误的 join key。",
        ],
    )

    doc.add_heading("8. 查看与导出结果", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "分析结果以结构化形式展示，而不是只返回一段文字。这样更利于复核、导出和后续汇报。"
    )
    set_paragraph_font(p)
    add_bullets(
        doc,
        [
            "Answers：按问题逐条展示最终回答。",
            "Key Metrics：展示关键指标。",
            "Tables：展示分析表格，页面展示会限制过大的表格以保持流畅。",
            "Visuals：展示图表；图表不是强制输出，只有在确实有助于理解时才生成。",
            "Findings / Insights：展示业务洞察或异常发现。",
            "Execution Details：展示执行相关信息，便于定位问题。",
        ],
    )
    doc.add_heading("8.1 导出规则", level=2)
    add_bullets(
        doc,
        [
            "点击 Export 可导出分析结果。",
            "当存在多个 result panel 时，可以选择导出全部结果或只导出某一个问题的结果。",
            "导出文件格式为 .xlsx，适合后续汇报、复核和留档。",
        ],
    )

    doc.add_heading("9. 大文件使用建议", level=1)
    add_bullets(
        doc,
        [
            "接近 1GB 的 Excel 文件可能需要更长导入和 profiling 时间。",
            "大型数据分析会先进行样本预检，再执行完整分析。",
            "尽量关闭无关大型程序，保证本机内存和磁盘空间充足。",
            "复杂多文件分析建议先明确关联键和分析粒度。",
            "如任务较重，优先提出单一清晰问题，再逐步追加分析。",
        ],
    )

    doc.add_heading("10. 常见问题与处理", level=1)
    faq = doc.add_table(rows=1, cols=3)
    faq.style = "Table Grid"
    set_cell_margins(faq)
    set_table_width(faq, [1.65, 2.25, 2.6])
    for idx, header in enumerate(["问题", "可能原因", "建议处理"]):
        cell = faq.rows[0].cells[idx]
        set_cell_shading(cell, BLUE_FILL)
        cell.paragraphs[0].add_run(header).bold = True
        set_paragraph_font(cell.paragraphs[0], size=9)
    rows = [
        ("无法导入文件", "文件类型不支持、超过 1GB、文件被占用或损坏", "确认文件为 .xlsx/.xls/.xlsm，关闭 Excel 后重试"),
        ("导入很慢", "文件较大或工作表较多", "等待后台处理完成，避免重复导入"),
        ("Please import a dataset first", "当前没有 Ready 状态数据集或未进入正确功能范围", "确认数据集已导入并在 Data Library 中高亮"),
        ("分析结果不符合预期", "问题过于宽泛，或多文件关联逻辑不明确", "补充字段、指标、关联键和期望输出"),
        ("代码需要修复", "AI 生成代码在本地执行时遇到字段、类型或关联问题", "等待自动修复；如多次失败，简化问题或明确列名"),
        ("清洗规则无法点击", "扫描后该规则不适用于当前数据", "无需处理；只有发现对应问题时规则才会启用"),
    ]
    for values in rows:
        cells = faq.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].paragraphs[0].add_run(value)
            set_paragraph_font(cells[idx].paragraphs[0], size=9)

    doc.add_heading("11. 最佳实践", level=1)
    add_bullets(
        doc,
        [
            "先检查数据集 overview，再提出分析问题。",
            "一次分析最多选择必要的数据集，不要无目的地全选。",
            "清洗前保留原始文件，清洗后抽样复核关键列。",
            "多问题分析时使用编号，例如“1. 统计销售额；2. 分析维修成本”。",
            "对正式汇报结果，建议查看 Python 代码和导出文件后再引用结论。",
            "如果字段名存在歧义，在问题中直接写出字段名。",
        ],
    )

    doc.add_heading("12. 简短术语说明", level=1)
    terms = doc.add_table(rows=1, cols=2)
    terms.style = "Table Grid"
    set_cell_margins(terms)
    set_table_width(terms, [1.65, 4.85])
    for idx, header in enumerate(["术语", "含义"]):
        cell = terms.rows[0].cells[idx]
        set_cell_shading(cell, BLUE_FILL)
        cell.paragraphs[0].add_run(header).bold = True
        set_paragraph_font(cell.paragraphs[0], size=9)
    for term, desc in [
        ("Dataset", "导入应用并可被清洗或分析的数据文件。"),
        ("Data Library", "全局数据资料库，用于查看、选择和删除数据集。"),
        ("Profiling / Scan", "对数据结构、质量问题和基础统计进行扫描。"),
        ("Dify", "用于生成分析计划和 Python 分析代码的 AI 工作流平台。"),
        ("Preflight", "在样本或代表性数据上预先检查代码是否可运行。"),
        ("Apply", "确认并执行当前 Python 分析代码。"),
        ("Result Panel", "按问题组织的结构化分析结果面板。"),
    ]:
        cells = terms.add_row().cells
        cells[0].paragraphs[0].add_run(term).bold = True
        cells[1].paragraphs[0].add_run(desc)
        set_paragraph_font(cells[0].paragraphs[0], size=9)
        set_paragraph_font(cells[1].paragraphs[0], size=9)

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(14)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run("End of Guide")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
