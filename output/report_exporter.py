"""
output/report_exporter.py
─────────────────────────
将报告导出为 Word 文档。

Sprint 3 🟡
"""

from __future__ import annotations
from pathlib import Path

# 如果未安装 python-docx，可选
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportExporter:
    """Word 报告导出器"""

    def __init__(self) -> None:
        if not HAS_DOCX:
            raise ImportError(
                "需要安装 python-docx: pip install python-docx"
            )

    def export(
        self,
        report_text: str,
        output_path: str,
        title: str = "数据分析报告",
        metadata: dict = None,
    ) -> None:
        """
        将报告导出为 Word 文档。
        
        Args:
            report_text: 报告文本
            output_path: 输出文件路径
            title: 文档标题
            metadata: 元数据（用户、日期等）
        """
        if not HAS_DOCX:
            raise RuntimeError("python-docx 未安装")

        doc = Document()

        # 标题
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = 1  # 居中

        # 元数据
        if metadata:
            doc.add_paragraph(
                f"生成时间: {metadata.get('timestamp', 'N/A')}"
            )
            doc.add_paragraph(
                f"文件: {metadata.get('file_name', 'N/A')}"
            )
            doc.add_paragraph(
                f"问题: {metadata.get('query', 'N/A')}"
            )
            doc.add_paragraph("")  # 空行

        # 报告内容
        doc.add_paragraph(report_text)

        # 保存
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)